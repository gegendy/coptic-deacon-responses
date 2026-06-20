from docx import Document
from docx.shared import Inches, Pt, Cm
from docx.enum.section import WD_ORIENT
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import nsdecls, qn
from docx.oxml import OxmlElement, parse_xml

# ============ FILE CONFIGURATION ============
DIVINE_LITURGY_INPUT = 'copticreader/divine-liturgy_deacon_responses_full.md'
DIVINE_LITURGY_OUTPUT = 'copticreader/divine-liturgy_deacon_responses.docx'

VESPER_MATINS_INPUT = 'copticreader/vesper-matins_deacon_responses_full.md'
VESPER_MATINS_OUTPUT = 'copticreader/vesper-matins_deacon_responses.docx'

COPTIC_FONT_NAME = 'Avva_Shenouda'

COPTIC_TO_AVVA_SHENOUDA = {
    'ϣ': '2',
    'ϥ': '4',
    'Ϧ': 'Q',
    'ϧ': 'q',
    'Ϩ': 'H',
    'ϩ': 'h',
    'ϫ': 'g',
    'Ϭ': 'S',
    'ϭ': 's',
    'ϯ': '5',
    'Ⲁ': 'A',
    'ⲁ': 'a',
    'Ⲃ': 'B',
    'ⲃ': 'b',
    'Ⲅ': 'J',
    'ⲅ': 'j',
    'ⲇ': 'd',
    'Ⲉ': 'E',
    'ⲉ': 'e',
    'Ⲍ': 'z',
    'ⲍ': 'z',
    'Ⲏ': '#',
    'ⲏ': '3',
    'Ⲑ': ')',
    'ⲑ': '0',
    'Ⲓ': 'I',
    'ⲓ': 'i',
    'Ⲕ': 'K',
    'ⲕ': 'k',
    'ⲗ': 'l',
    'Ⲙ': 'M',
    'ⲙ': 'm',
    'Ⲛ': 'N',
    'ⲛ': 'n',
    'Ⲝ': '7',
    'ⲝ': '7',
    'Ⲟ': 'O',
    'ⲟ': 'o',
    'Ⲡ': 'P',
    'ⲡ': 'p',
    'ⲣ': 'r',
    'Ⲥ': 'C',
    'ⲥ': 'c',
    'Ⲧ': 'T',
    'ⲧ': 't',
    'Ⲩ': 'V',
    'ⲩ': 'v',
    'Ⲫ': 'F',
    'ⲫ': 'f',
    'Ⲭ': 'X',
    'ⲭ': 'x',
    'Ⲯ': 'Y',
    'ⲯ': 'y',
    'ⲱ': 'w',
}
# ============================================

def set_run_font(run, font_name, size=None, bold=None):
    """Apply a font across all Word script slots for the run."""
    run.font.name = font_name
    r_pr = run._element.get_or_add_rPr()
    r_fonts = r_pr.rFonts
    if r_fonts is None:
        r_fonts = OxmlElement('w:rFonts')
        r_pr.append(r_fonts)
    for attr in ('ascii', 'hAnsi', 'eastAsia', 'cs'):
        r_fonts.set(qn(f'w:{attr}'), font_name)
    if size is not None:
        run.font.size = size
    if bold is not None:
        run.font.bold = bold

def convert_coptic_text(text):
    """Convert Unicode Coptic text into legacy Avva Shenouda encoding."""
    converted = []
    index = 0
    while index < len(text):
        char = text[index]
        next_char = text[index + 1] if index + 1 < len(text) else None

        if char in ('\u0300', '\u0305'):
            raise ValueError(f'Unexpected standalone combining mark in Coptic text: {char!r}')

        if char in COPTIC_TO_AVVA_SHENOUDA:
            legacy_char = COPTIC_TO_AVVA_SHENOUDA[char]
            if next_char == '\u0305':
                converted.append('=')
                converted.append(legacy_char)
                index += 2
                continue
            if next_char == '\u0300':
                converted.append('`')
                converted.append(legacy_char)
                index += 2
                continue
            converted.append(legacy_char)
        else:
            codepoint = ord(char)
            if 0x2C80 <= codepoint <= 0x2CFF or 0x03E2 <= codepoint <= 0x03EF:
                raise ValueError(f'Unsupported Coptic character for Avva Shenouda conversion: {char!r}')
            converted.append(char)

        index += 1

    return ''.join(converted)

def add_coptic_runs(paragraph, text, size):
    """Render Coptic text into a paragraph using multiple runs.

    Coptic letters and their combining marks are emitted in the Avva Shenouda
    legacy font, while every other character (spaces, parentheses, commas,
    colons, periods, Latin letters, etc.) is emitted in the paragraph's default
    font. The legacy font maps ASCII punctuation to Coptic glyphs (for example
    ')' renders as theta and '(' / ',' render blank), so such characters must
    not be placed in the legacy font.
    """
    avva_buffer = []
    plain_buffer = []

    def flush_avva():
        if avva_buffer:
            run = paragraph.add_run(''.join(avva_buffer))
            set_run_font(run, COPTIC_FONT_NAME, size=size)
            avva_buffer.clear()

    def flush_plain():
        if plain_buffer:
            run = paragraph.add_run(''.join(plain_buffer))
            run.font.size = size
            plain_buffer.clear()

    index = 0
    while index < len(text):
        char = text[index]
        next_char = text[index + 1] if index + 1 < len(text) else None

        if char in ('\u0300', '\u0305'):
            raise ValueError(f'Unexpected standalone combining mark in Coptic text: {char!r}')

        if char in COPTIC_TO_AVVA_SHENOUDA:
            flush_plain()
            legacy_char = COPTIC_TO_AVVA_SHENOUDA[char]
            if next_char == '\u0305':
                avva_buffer.append('=')
                avva_buffer.append(legacy_char)
                index += 2
                continue
            if next_char == '\u0300':
                avva_buffer.append('`')
                avva_buffer.append(legacy_char)
                index += 2
                continue
            avva_buffer.append(legacy_char)
        else:
            codepoint = ord(char)
            if 0x2C80 <= codepoint <= 0x2CFF or 0x03E2 <= codepoint <= 0x03EF:
                raise ValueError(f'Unsupported Coptic character for Avva Shenouda conversion: {char!r}')
            flush_avva()
            plain_buffer.append(char)

        index += 1

    flush_avva()
    flush_plain()

def set_cell_shading(cell, color):
    """Set cell background color"""
    shading = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{color}"/>')
    cell._tc.get_or_add_tcPr().append(shading)

def prevent_row_break(row):
    """Prevent row from breaking across pages"""
    tr = row._tr
    trPr = tr.get_or_add_trPr()
    for child in trPr:
        if child.tag == qn('w:cantSplit'):
            trPr.remove(child)
    cant_split = parse_xml(f'<w:cantSplit {nsdecls("w")} w:val="true"/>')
    trPr.append(cant_split)

def create_table_from_data(doc, title, rows, is_contents=False):
    """Create a formatted table"""
    if title:
        heading = doc.add_heading(title, level=1)
        heading.alignment = WD_ALIGN_PARAGRAPH.CENTER
        for run in heading.runs:
            run.font.size = Pt(11)
            run.font.bold = True

    num_cols = len(rows[0]) if rows else 4

    table = doc.add_table(rows=1, cols=num_cols)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False

    header_cells = table.rows[0].cells
    if num_cols == 5:
        headers = ['Response', 'English', 'Coptic', 'Arabic', 'Transliteration']
    else:
        headers = ['English', 'Coptic', 'Arabic', 'Transliteration']
    coptic_idx = 2 if num_cols == 5 else 1

    for i, header in enumerate(headers):
        header_cells[i].text = header
        header_cells[i].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        for para in header_cells[i].paragraphs:
            para.paragraph_format.space_before = Pt(0)
            para.paragraph_format.space_after = Pt(0)
        for run in header_cells[i].paragraphs[0].runs:
            run.font.bold = True
            run.font.size = Pt(5)
        set_cell_shading(header_cells[i], "D9E2F3")

    prevent_row_break(table.rows[0])

    for row_data in rows:
        row = table.add_row()
        prevent_row_break(row)
        for i, cell_text in enumerate(row_data):
            cell = row.cells[i]
            if i == coptic_idx:
                para = cell.paragraphs[0]
                add_coptic_runs(para, cell_text, Pt(5))
            else:
                cell.text = cell_text
                para = cell.paragraphs[0]
            para.paragraph_format.space_before = Pt(0)
            para.paragraph_format.space_after = Pt(0)
            arabic_idx = 3 if num_cols == 5 else 2
            if i == arabic_idx:
                para.alignment = WD_ALIGN_PARAGRAPH.RIGHT
            if i != coptic_idx:
                for run in para.runs:
                    run.font.size = Pt(5)

    if num_cols == 5:
        widths = [Cm(5.5)] * 5
    else:
        widths = [Cm(6.88)] * 4

    for row in table.rows:
        for i, cell in enumerate(row.cells):
            cell.width = widths[i]
    for i, column in enumerate(table.columns):
        column.width = widths[i]

def parse_markdown_table(content):
    """Parse markdown table and return rows"""
    lines = content.strip().split('\n')
    rows = []
    for line in lines:
        if line.startswith('|') and not line.startswith('|--'):
            if '---' in line:
                continue
            cells = [cell.strip() for cell in line.split('|')[1:-1]]
            if len(cells) == 5:
                if cells[0] == 'Response Name' or (cells[1] == '' and cells[2] == ''):
                    continue
                rows.append(cells)
            elif len(cells) == 4 and cells[0] and cells[0] != 'English':
                rows.append(cells)
    return rows

# Read the markdown file
with open(DIVINE_LITURGY_INPUT, 'r', encoding='utf-8') as f:
    full_md = f.read()

full_rows = parse_markdown_table(full_md)

doc = Document()

section = doc.sections[0]
section.orientation = WD_ORIENT.LANDSCAPE
section.page_width = Inches(11)
section.page_height = Inches(8.5)
section.top_margin = Cm(0.2)
section.bottom_margin = Cm(0.2)
section.left_margin = Cm(0.2)
section.right_margin = Cm(0.2)
section.header_distance = Cm(0)
section.footer_distance = Cm(0)

title = doc.add_heading("Divine Liturgy Deacon Responses", level=0)
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
title.paragraph_format.space_before = Pt(0)
title.paragraph_format.space_after = Pt(2)
for run in title.runs:
    run.font.size = Pt(12)

create_table_from_data(doc, "", full_rows, is_contents=False)

doc.save(DIVINE_LITURGY_OUTPUT)
print(f"Word document created: {DIVINE_LITURGY_OUTPUT}")

# --- Now generate Vespers/Matins document ---

with open(VESPER_MATINS_INPUT, 'r', encoding='utf-8') as f:
    vesper_md = f.read()

vesper_rows = parse_markdown_table(vesper_md)

doc2 = Document()

section2 = doc2.sections[0]
section2.orientation = WD_ORIENT.LANDSCAPE
section2.page_width = Inches(11)
section2.page_height = Inches(8.5)
section2.top_margin = Cm(0.2)
section2.bottom_margin = Cm(0.2)
section2.left_margin = Cm(0.2)
section2.right_margin = Cm(0.2)
section2.header_distance = Cm(0)
section2.footer_distance = Cm(0)

title2 = doc2.add_heading("Vespers & Matins Deacon Responses", level=0)
title2.alignment = WD_ALIGN_PARAGRAPH.CENTER
title2.paragraph_format.space_before = Pt(0)
title2.paragraph_format.space_after = Pt(2)
for run in title2.runs:
    run.font.size = Pt(12)

create_table_from_data(doc2, "", vesper_rows, is_contents=False)

doc2.save(VESPER_MATINS_OUTPUT)
print(f"Word document created: {VESPER_MATINS_OUTPUT}")
